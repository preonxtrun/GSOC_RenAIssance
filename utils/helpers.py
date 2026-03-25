"""
Utility helpers: .docx parsing, logging, device selection.
"""
import logging
import torch
from pathlib import Path
from docx import Document


def setup_logger(name: str = "renaissance-htr", level=logging.INFO) -> logging.Logger:
    """Create a formatted logger."""
    logger = logging.getLogger(name)
    if not logger.handlers:
        handler = logging.StreamHandler()
        formatter = logging.Formatter(
            "[%(asctime)s] %(levelname)s - %(name)s - %(message)s",
            datefmt="%H:%M:%S",
        )
        handler.setFormatter(formatter)
        logger.addHandler(handler)
    logger.setLevel(level)
    return logger


def get_device() -> torch.device:
    """Get the best available device."""
    if torch.cuda.is_available():
        device = torch.device("cuda")
    elif hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
        device = torch.device("mps")
    else:
        device = torch.device("cpu")
    return device


def parse_docx(docx_path: str | Path) -> str:
    """
    Extract all text from a .docx transcription file.
    Returns the full text as a single string with paragraph breaks.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Transcription file not found: {docx_path}")

    doc = Document(str(docx_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Skip empty paragraphs
            paragraphs.append(text)

    return "\n".join(paragraphs)


def parse_docx_paragraphs(docx_path: str | Path) -> list[str]:
    """
    Extract text from a .docx file as a list of non-empty paragraphs.
    Useful for paragraph-level alignment with page images.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Transcription file not found: {docx_path}")

    doc = Document(str(docx_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return paragraphs
